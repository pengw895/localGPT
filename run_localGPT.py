import os
import logging
import click
import torch
import utils
from langchain.chains import RetrievalQA
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.llms import HuggingFacePipeline
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler  # for streaming response
from langchain.callbacks.manager import CallbackManager
from docx import Document
import re

callback_manager = CallbackManager([StreamingStdOutCallbackHandler()])

from prompt_template_utils import get_prompt_template
from utils import get_embeddings

# from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.vectorstores import Chroma
from transformers import (
    GenerationConfig,
    pipeline,
)

from load_models import (
    load_quantized_model_awq,
    load_quantized_model_gguf_ggml,
    load_quantized_model_qptq,
    load_full_model,
)

from constants import (
    EMBEDDING_MODEL_NAME,
    PERSIST_DIRECTORY,
    MODEL_ID,
    MODEL_BASENAME,
    MAX_NEW_TOKENS,
    MODELS_PATH,
    CHROMA_SETTINGS,    
)


def load_model(device_type, model_id, model_basename=None, LOGGING=logging):
    """
    Select a model for text generation using the HuggingFace library.
    If you are running this for the first time, it will download a model for you.
    subsequent runs will use the model from the disk.

    Args:
        device_type (str): Type of device to use, e.g., "cuda" for GPU or "cpu" for CPU.
        model_id (str): Identifier of the model to load from HuggingFace's model hub.
        model_basename (str, optional): Basename of the model if using quantized models.
            Defaults to None.

    Returns:
        HuggingFacePipeline: A pipeline object for text generation using the loaded model.

    Raises:
        ValueError: If an unsupported model or device type is provided.
    """
    logging.info(f"Loading Model: {model_id}, on: {device_type}")
    logging.info("This action can take a few minutes!")
    
    if model_basename is not None:
        if ".gguf" in model_basename.lower():
            llm = load_quantized_model_gguf_ggml(model_id, model_basename, device_type, LOGGING)
            return llm
        elif ".ggml" in model_basename.lower():
            model, tokenizer = load_quantized_model_gguf_ggml(model_id, model_basename, device_type, LOGGING)
        elif ".awq" in model_basename.lower():
            model, tokenizer = load_quantized_model_awq(model_id, LOGGING)
        else:
            model, tokenizer = load_quantized_model_qptq(model_id, model_basename, device_type, LOGGING)
    else:
        model, tokenizer = load_full_model(model_id, model_basename, device_type, LOGGING)

    # Load configuration from the model to avoid warnings
    generation_config = GenerationConfig.from_pretrained(model_id)
    # see here for details:
    # https://huggingface.co/docs/transformers/
    # main_classes/text_generation#transformers.GenerationConfig.from_pretrained.returns

    # Create a pipeline for text generation
    if device_type == "hpu":
        from gaudi_utils.pipeline import GaudiTextGenerationPipeline

        pipe = GaudiTextGenerationPipeline(
            model_name_or_path=model_id,
            max_new_tokens=1000,
            temperature=0.2,
            top_p=0.95,
            repetition_penalty=1.15,
            do_sample=True,
            max_padding_length=5000,
        )
        pipe.compile_graph()
    else:
        pipe = pipeline(
        "text-generation",
        model=model,
        tokenizer=tokenizer,
        max_length=MAX_NEW_TOKENS,
        temperature=0.2,
        # top_p=0.95,
        repetition_penalty=1.15,
        generation_config=generation_config,
    )

    local_llm = HuggingFacePipeline(pipeline=pipe)
    logging.info("Local LLM Loaded")

    return local_llm


def retrieval_qa_pipline(device_type, use_history, promptTemplate_type="llama"):
    """
    Initializes and returns a retrieval-based Question Answering (QA) pipeline.

    This function sets up a QA system that retrieves relevant information using embeddings
    from the HuggingFace library. It then answers questions based on the retrieved information.

    Parameters:
    - device_type (str): Specifies the type of device where the model will run, e.g., 'cpu', 'cuda', etc.
    - use_history (bool): Flag to determine whether to use chat history or not.

    Returns:
    - RetrievalQA: An initialized retrieval-based QA system.

    Notes:
    - The function uses embeddings from the HuggingFace library, either instruction-based or regular.
    - The Chroma class is used to load a vector store containing pre-computed embeddings.
    - The retriever fetches relevant documents or data based on a query.
    - The prompt and memory, obtained from the `get_prompt_template` function, might be used in the QA system.
    - The model is loaded onto the specified device using its ID and basename.
    - The QA system retrieves relevant documents using the retriever and then answers questions based on those documents.
    """

    """
    (1) Chooses an appropriate langchain library based on the enbedding model name.  Matching code is contained within ingest.py.

    (2) Provides additional arguments for instructor and BGE models to improve results, pursuant to the instructions contained on
    their respective huggingface repository, project page or github repository.
    """
    if device_type == "hpu":
        from gaudi_utils.embeddings import load_embeddings

        embeddings = load_embeddings()
    else:
        embeddings = get_embeddings(device_type)

    logging.info(f"Loaded embeddings from {EMBEDDING_MODEL_NAME}")

    # load the vectorstore
    db = Chroma(persist_directory=PERSIST_DIRECTORY, embedding_function=embeddings, client_settings=CHROMA_SETTINGS)
    retriever = db.as_retriever()

    # get the prompt template and memory if set by the user.
    prompt, memory = get_prompt_template(promptTemplate_type=promptTemplate_type, history=use_history)

    # load the llm pipeline
    llm = load_model(device_type, model_id=MODEL_ID, model_basename=MODEL_BASENAME, LOGGING=logging)

    if use_history:
        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",  # try other chains types as well. refine, map_reduce, map_rerank
            retriever=retriever,
            return_source_documents=True,  # verbose=True,
            callbacks=callback_manager,
            chain_type_kwargs={"prompt": prompt, "memory": memory},
        )
    else:
        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",  # try other chains types as well. refine, map_reduce, map_rerank
            retriever=retriever,
            return_source_documents=True,  # verbose=True,
            callbacks=callback_manager,
            chain_type_kwargs={
                "prompt": prompt,
            },
        )

    return qa


def is_generic_content(text, qa):
    """
    Use the LLM to determine if a piece of text is generic/placeholder content.
    Returns True if the text appears to be a placeholder that needs filling.
    """
    if not text.strip():
        return False
        
    # Query the LLM to analyze the text
    query = f"""Analyze this text and determine if it's a generic placeholder that needs to be filled in with specific information.
    Consider it generic if it:
    1. Contains words like 'your', 'the', 'a', 'an' followed by a noun without specific details
    2. Uses placeholder-like language (e.g., 'enter your', 'fill in', 'specify')
    3. Is too vague or general to be meaningful content
    4. Looks like it's waiting for specific information to be inserted
    
    Text to analyze: "{text}"
    
    Respond with only 'GENERIC' if it's a placeholder, or 'SPECIFIC' if it contains meaningful content.
    """
    
    res = qa(query)
    answer = res["result"].strip().upper()
    return "GENERIC" in answer

def find_generic_content(doc, qa):
    """
    Find all generic/placeholder content in a Word document using LLM analysis.
    Prints each piece of generic content as it's found.
    Returns a list of tuples (text, location) where location describes where in the document the text was found.
    """
    generic_content = []
    
    print("\nAnalyzing document for generic content...")
    
    # Check paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            if is_generic_content(paragraph.text, qa):
                location = f"Paragraph {i+1}"
                print(f"\nFound generic content in {location}:")
                print(f"Content: {paragraph.text}")
                generic_content.append((paragraph.text, location))
    
    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():
                        if is_generic_content(paragraph.text, qa):
                            location = f"Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}, Paragraph {para_idx+1}"
                            print(f"\nFound generic content in {location}:")
                            print(f"Content: {paragraph.text}")
                            generic_content.append((paragraph.text, location))
    
    return generic_content

def fill_generic_content(doc, generic_content, qa):
    """
    For each piece of generic content, analyze the template structure and reformat the CI6782 document accordingly.
    First handles standard replacements, then uses module requirement verification content for the MR table.
    Returns list of unfilled generic content.
    """
    unfilled_content = []
    
    # First, get an overview of the template structure
    template_structure = []
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            template_structure.append(f"Paragraph {i+1}: {paragraph.text}")
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():
                        template_structure.append(
                            f"Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}, Paragraph {para_idx+1}: {paragraph.text}"
                        )
    
    # Create a comprehensive prompt for reformatting
    template_overview = "\n".join(template_structure)
    
    # First, get the module requirement verification content with a more specific query
    mr_query = """Search the CI6782-140-001 document for the Module Requirements section.
    Look for content that describes:
    1. Each requirement's ID (e.g., CI6782-XXX, etc.)
    2. The specific requirement text
    3. How the requirement was verified
    4. The verification results
    
    Format each requirement as:
    Requirement ID: [ID]
    Description: [requirement text]
    Verification Method: [method used]
    Result: [verification result]
    
    Only return actual content from the document, no made-up or placeholder content."""
    
    print("\nFetching module requirements content...")
    mr_res = qa(mr_query)
    mr_content = mr_res["result"]
    
    # Verify we got meaningful content
    if not mr_content or "MS MS" in mr_content or len(mr_content.strip()) < 50:
        print("Warning: Could not retrieve meaningful module requirements content.")
        mr_content = "Module requirements content could not be retrieved from the document."
    
    reformat_prompt = f"""You are tasked with reformatting the CI6782-140-001 document according to this template structure.

Template Structure:
{template_overview}

Available Module Requirements Content:
{mr_content}

Instructions:
1. First, replace these standard fields:
   - Replace 'Document Title' with 'Cypress Privacy & Security Requirements Verification by Analysis'
   - Replace 'Document Number' with 'CI6782-140-001'
   - Replace 'Project Name' with 'Cypress'
   - Replace 'Project Number' with 'CI6782'
   - Replace 'Document Date' with '2024-04-06'

2. Then rewrite the CI6782-140-001 document to match the template structure:
   - Use the provided module requirements content to fill in the MR table
   - Maintain the template's formatting and structure
   - Ensure all sections are properly connected
   - Keep the document's professional tone and style

Important:
- Only use actual content from the CI6782-140-001 document
- Do not make up or generate placeholder content
- If you cannot find specific content, indicate this clearly
- Maintain the exact formatting of the template

Return the reformatted content for this specific location."""

    for text, location in generic_content:
        # Query the RAG system for reformatted content
        query = f"{reformat_prompt}\n\nLocation to fill: {location}\nCurrent generic content: {text}"
        
        print(f"\nProcessing location: {location}")
        res = qa(query)
        answer = res["result"]
        
        # Verify we got meaningful content
        if not answer or "MS MS" in answer or len(answer.strip()) < 10:
            print(f"Warning: Could not get meaningful content for {location}")
            unfilled_content.append((text, location))
            continue
        
        if "I don't know" not in answer and "cannot find" not in answer.lower():
            print(f"\nReformatting content for {location}:")
            print(f"Original template text: {text}")
            print(f"Reformatted content: {answer}")
            
            # Replace the content in the document
            if "Table" in location:
                # Parse location to get table, row, cell, and paragraph indices
                parts = location.split(", ")
                table_idx = int(parts[0].split()[1]) - 1
                row_idx = int(parts[1].split()[1]) - 1
                cell_idx = int(parts[2].split()[1]) - 1
                para_idx = int(parts[3].split()[1]) - 1
                
                doc.tables[table_idx].rows[row_idx].cells[cell_idx].paragraphs[para_idx].text = answer
            else:
                # Parse paragraph index
                para_idx = int(location.split()[1]) - 1
                doc.paragraphs[para_idx].text = answer
        else:
            print(f"\nCould not find appropriate content for {location}:")
            print(f"Generic content: {text}")
            unfilled_content.append((text, location))
    
    return unfilled_content

def process_template(template_path, qa):
    """
    Process a Word template document:
    1. Find all generic/placeholder content using LLM
    2. For each generic content, query the RAG system for specific information
    3. Fill in the template with found values
    4. Return list of unfilled generic content
    """
    doc = Document(template_path)
    
    print("\nAnalyzing document for generic content...")
    generic_content = find_generic_content(doc, qa)
    
    print("\nFound generic content:")
    for text, location in generic_content:
        print(f"\nLocation: {location}")
        print(f"Content: {text}")
    
    print("\nAttempting to fill generic content with specific information...")
    unfilled_content = fill_generic_content(doc, generic_content, qa)
    
    # Save the filled template
    output_path = "filled-in-template.docx"
    doc.save(output_path)
    print(f"\nSaved filled template to {output_path}")
    
    return unfilled_content

# chose device typ to run on as well as to show source documents.
@click.command()
@click.option(
    "--device_type",
    default="cuda" if torch.cuda.is_available() else "cpu",
    type=click.Choice(
        [
            "cpu",
            "cuda",
            "ipu",
            "xpu",
            "mkldnn",
            "opengl",
            "opencl",
            "ideep",
            "hip",
            "ve",
            "fpga",
            "ort",
            "xla",
            "lazy",
            "vulkan",
            "mps",
            "meta",
            "hpu",
            "mtia",
        ],
    ),
    help="Device to run on. (Default is cuda)",
)
@click.option(
    "--show_sources",
    "-s",
    is_flag=True,
    help="Show sources along with answers (Default is False)",
)
@click.option(
    "--use_history",
    "-h",
    is_flag=True,
    help="Use history (Default is False)",
)
@click.option(
    "--model_type",
    default="llama3",
    type=click.Choice(
        ["llama3", "llama", "mistral", "non_llama"],
    ),
    help="model type, llama3, llama, mistral or non_llama",
)
@click.option(
    "--save_qa",
    is_flag=True,
    help="whether to save Q&A pairs to a CSV file (Default is False)",
)
@click.option(
    "--template_path",
    default="template.docx",
    help="Path to the Word template file (Default is template.docx)",
)
def main(device_type, show_sources, use_history, model_type, save_qa, template_path):
    """
    Process a Word template document using RAG to fill in placeholders.
    """
    logging.info(f"Running on: {device_type}")
    logging.info(f"Display Source Documents set to: {show_sources}")
    logging.info(f"Use history set to: {use_history}")

    # check if models directory do not exist, create a new one and store models here.
    if not os.path.exists(MODELS_PATH):
        os.mkdir(MODELS_PATH)

    # Initialize the QA system
    qa = retrieval_qa_pipline(device_type, use_history, promptTemplate_type=model_type)
    
    # Process the template
    if not os.path.exists(template_path):
        print(f"Error: Template file {template_path} not found!")
        return
        
    unfilled_content = process_template(template_path, qa)
    
    if unfilled_content:
        print("\nThe following placeholders could not be filled:")
        for text, location in unfilled_content:
            print(f"- {text} (Location: {location})")
    else:
        print("\nAll placeholders were successfully filled!")


if __name__ == "__main__":
    logging.basicConfig(
        format="%(asctime)s - %(levelname)s - %(filename)s:%(lineno)s - %(message)s", level=logging.INFO
    )
    main()
